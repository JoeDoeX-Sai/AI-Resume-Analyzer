"""
Resume Parser - Extracts raw text from PDF, DOCX, and TXT files.
"""
import re
from pathlib import Path


class ResumeParser:
    """Handles text extraction from multiple resume formats."""

    def parse(self, filepath: str) -> str:
        ext = Path(filepath).suffix.lower()
        if ext == ".pdf":
            return self._parse_pdf(filepath)
        elif ext == ".docx":
            return self._parse_docx(filepath)
        elif ext == ".txt":
            return self._parse_txt(filepath)
        else:
            raise ValueError(f"Unsupported file format: {ext}")

    def _parse_pdf(self, filepath: str) -> str:
        from pdfminer.high_level import extract_text
        text = extract_text(filepath)
        return self._clean_text(text)

    def _parse_docx(self, filepath: str) -> str:
        from docx import Document
        doc = Document(filepath)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return self._clean_text("\n".join(paragraphs))

    def _parse_txt(self, filepath: str) -> str:
        with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
            return self._clean_text(f.read())

    def _clean_text(self, text: str) -> str:
        if not text:
            return ""
        # Normalize whitespace
        text = re.sub(r"\r\n|\r", "\n", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        text = re.sub(r"[ \t]{2,}", " ", text)
        # Remove non-printable characters
        text = re.sub(r"[^\x20-\x7E\n]", " ", text)
        return text.strip()
